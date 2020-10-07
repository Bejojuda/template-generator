import os
import re
from io import BytesIO

from django.core.files.base import ContentFile
from django.http import HttpResponse
from docx import Document
from rest_framework.exceptions import ValidationError

from variable.models import Variable


def search_variables(document):
    """
    Search and return all of the variables in the document.
    A variable has this format: {{variable_name}}
    """
    variables = []
    for p in document.paragraphs:
        variables += re.findall(r"\{\{(.[^{}]*?)\}\}", p.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    variables += re.findall(r"\{\{(.[^{}]*?)\}\}", p.text)
    variables = list(dict.fromkeys(variables))
    return variables


def search_optional_variables(document):
    """
    Search and return all of the optional variables in the document.
    An optional variable has this format: [[variable_name]]
    """
    optional_variables = []
    for p in document.paragraphs:
        optional_variables += re.findall(r"\[\[(.[^{}]*?)\]\]", p.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    optional_variables += re.findall(r"\[\[(.[^{}]*?)\]\]", p.text)
    optional_variables = list(dict.fromkeys(optional_variables))
    return optional_variables


def create_variables(template, variables_str=[], optional_variables_str=[]):
    """
    Receives a list of variables and a template, and then creates the variable objects and the relationship
    with said template object
    """
    for v in variables_str:
        variable = Variable.objects.create(name=v, template=template)
        template.variables.add(variable)

    for ov in optional_variables_str:
        optional_variable = Variable.objects.create(name=ov, template=template, optional=True)
        template.variables.add(optional_variable)


def replace_variables(doc, var={}, opt={}):
    """
    Receives a variables dictionary and replaces the variables inside the document 'doc' to then return a new document
    """
    variables = {}
    optional = {}
    if not var:
        return None

    for v in var:
        variables['{{'+v+'}}'] = var[v]
    for o in opt:
        variables['[['+o+']]'] = opt[o]
        optional['[['+o+']]'] = True
    paragraphs = list(doc.paragraphs)
    # Iterates the document tables to get the text inside each cell
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        for search_text in variables:
            replace_text = variables[search_text]
            if search_text in p.text:
                for _ in range(p.text.count(search_text)):
                    inline = p.runs
                    # Replace strings and retain the same style.
                    # The text to be replaced can be split over several runs so
                    # search through, identify which runs need to have text replaced
                    # then replace the text in those identified
                    started = False
                    search_index = 0
                    # found_runs is a list of (inline index, index of match, length of match)
                    found_runs = list()
                    found_all = False
                    replace_done = False
                    for i in range(len(inline)):

                        # case 1: found in single run so short circuit the replace
                        if search_text in inline[i].text and not started:
                            found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                            text = inline[i].text.replace(search_text, str(replace_text))
                            inline[i].text = text
                            replace_done = True
                            found_all = True
                            break

                        if search_text[search_index] not in inline[i].text and not started:
                            # keep looking ...
                            continue

                        # case 2: search for partial text, find first run
                        if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                            # check sequence
                            start_index = inline[i].text.find(search_text[search_index])
                            check_length = len(inline[i].text)
                            for text_index in range(start_index, check_length):
                                if inline[i].text[text_index] != search_text[search_index]:
                                    # no match so must be false positive
                                    break
                            if search_index == 0:
                                started = True
                            chars_found = check_length - start_index
                            search_index += chars_found
                            found_runs.append((i, start_index, chars_found))
                            if search_index != len(search_text):
                                continue
                            else:
                                # found all chars in search_text
                                found_all = True
                                break

                        # case 2: search for partial text, find subsequent run
                        if search_text[search_index] in inline[i].text and started and not found_all:
                            # check sequence
                            chars_found = 0
                            check_length = len(inline[i].text)
                            for text_index in range(0, check_length):
                                if inline[i].text[text_index] == search_text[search_index]:
                                    search_index += 1
                                    chars_found += 1
                                else:
                                    break
                            # no match so must be end
                            found_runs.append((i, 0, chars_found))
                            if search_index == len(search_text):
                                found_all = True
                                break

                    if found_all and not replace_done:
                        for i, item in enumerate(found_runs):
                            index, start, length = [t for t in item]
                            if i == 0:
                                text = inline[index].text.replace(inline[index].text[start:start + length], str(replace_text))
                                inline[index].text = text
                            else:
                                text = inline[index].text.replace(inline[index].text[start:start + length], '')
                                inline[index].text = text

    return doc


def generate_document(data, template):
    """
    When a PUT is made with variables, a new .docx is generated with said variables added to the document.
    The document is returned to be downloaded by the user.
    """
    doc = Document(template.document)

    sent_variables = data['sent_variables']
    optional_variables = {}

    template_variables = Variable.objects.filter(template_id=template.uuid, optional=False)
    check_required_variables(sent_variables, template_variables)

    if 'optional_variables' in data:
        optional_variables = data['optional_variables']

    document = replace_variables(doc, sent_variables, optional_variables)

    f = BytesIO()
    document.save(f)
    response = HttpResponse(f.getvalue(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="%s"' % template.name

    return response


def get_upload_path(instance, filename):
    """
    Returns custom file path that has the creation date in the containing folder
    """
    return os.path.join(
      "documents", "date_%s" % instance.create_date, filename)


def check_required_variables(sent_v, template_v):
    """
    Checks if the sent variables are the same as the variables from the template
    """
    template_variables = [tv.name for tv in template_v]
    sent_variables = [sv for sv in sent_v]
    if not set(template_variables) == set(sent_variables):
        raise ValidationError({"error": "Sent variables are not correct"})


def rename_file(new_name, template, templates):
    """
    Changes the name of the file referred in the Template document FileField to new_name
    """
    pattern = re.compile('^[^<>:;,.?"*| /]+$')
    if not pattern.match(new_name):
        raise ValidationError({
            "error": "File name is incorrect"
        })

    for t in templates:
        if t.filename() == new_name:
            raise ValidationError({
                "error": "Document filename already exists"
            })

    old_path = template.document.path
    new_path = old_path.replace(template.document.path, new_name)

    with open(template.document.path, "rb") as f:
        with ContentFile(f.read()) as doc_content:
            template.document.save(new_path, doc_content)
            template.save()

    os.remove(old_path)

    return new_name
